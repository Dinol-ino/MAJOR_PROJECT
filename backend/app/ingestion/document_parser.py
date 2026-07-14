from __future__ import annotations

import mimetypes
import os
import tempfile
from pathlib import Path
from typing import Any

from app.config import settings
from app.ingestion.types import ParsedDocument, ParsedPage

try:
    import fitz
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


class UnsupportedDocumentError(ValueError):
    pass


class DocumentParser:
    def __init__(
        self,
        max_size_mb: int | None = None,
        max_pages: int | None = None,
        ocr_enabled: bool | None = None,
        ocr_min_text_threshold: float | None = None,
    ):
        self.max_size_mb = max_size_mb or settings.MAX_FILE_SIZE_MB
        self.max_pages = max_pages or settings.MAX_FILE_PAGES
        self.ocr_enabled = settings.OCR_ENABLED if ocr_enabled is None else ocr_enabled
        self.ocr_min_text_threshold = (
            settings.OCR_MIN_TEXT_THRESHOLD
            if ocr_min_text_threshold is None
            else ocr_min_text_threshold
        )
        self._ocr_engine: Any | None = None

    def parse(self, file_path: str | Path, filename: str | None = None) -> ParsedDocument:
        path = Path(file_path)
        source_name = filename or path.name
        extension = path.suffix.lower()

        self._validate_file(path, extension)

        if extension == ".pdf":
            return self._parse_pdf(path, source_name)
        if extension == ".txt":
            return self._parse_txt(path, source_name)
        if extension == ".docx":
            return self._parse_docx(path, source_name)

        raise UnsupportedDocumentError(f"Unsupported file type: {extension}")

    def _validate_file(self, path: Path, extension: str) -> None:
        if extension not in SUPPORTED_EXTENSIONS:
            raise UnsupportedDocumentError(
                "Supported document types are PDF, TXT, and DOCX."
            )

        file_size_mb = os.path.getsize(path) / (1024 * 1024)
        if file_size_mb > self.max_size_mb:
            raise ValueError(f"File exceeds size limit of {self.max_size_mb}MB.")

    def _parse_pdf(self, path: Path, filename: str) -> ParsedDocument:
        if fitz is not None:
            return self._parse_pdf_with_pymupdf(path, filename)
        if pdfplumber is not None:
            return self._parse_pdf_with_pdfplumber(path, filename)
        raise ImportError("Install PyMuPDF or pdfplumber to parse PDF files.")

    def _parse_pdf_with_pymupdf(self, path: Path, filename: str) -> ParsedDocument:
        pages: list[ParsedPage] = []
        with fitz.open(path) as document:
            if document.page_count > self.max_pages:
                raise ValueError(f"File exceeds page limit of {self.max_pages} pages.")

            for index, page in enumerate(document, start=1):
                text = page.get_text("text") or ""
                ocr_used = False
                if self._needs_ocr(text) and self.ocr_enabled:
                    ocr_text = self._try_ocr_page(page)
                    if ocr_text:
                        text = ocr_text
                        ocr_used = True
                pages.append(
                    ParsedPage(
                        page_number=index,
                        text=text.strip(),
                        ocr_used=ocr_used,
                    )
                )

        return ParsedDocument(
            filename=filename,
            mime_type="application/pdf",
            pages=pages,
            metadata={"parser": "pymupdf"},
        )

    def _parse_pdf_with_pdfplumber(self, path: Path, filename: str) -> ParsedDocument:
        pages: list[ParsedPage] = []
        with pdfplumber.open(path) as document:
            if len(document.pages) > self.max_pages:
                raise ValueError(f"File exceeds page limit of {self.max_pages} pages.")

            for index, page in enumerate(document.pages, start=1):
                text = page.extract_text() or ""
                pages.append(ParsedPage(page_number=index, text=text.strip()))

        return ParsedDocument(
            filename=filename,
            mime_type="application/pdf",
            pages=pages,
            metadata={"parser": "pdfplumber"},
        )

    def _parse_txt(self, path: Path, filename: str) -> ParsedDocument:
        text = path.read_text(encoding="utf-8")
        return ParsedDocument(
            filename=filename,
            mime_type="text/plain",
            pages=[ParsedPage(page_number=1, text=text.strip())],
            metadata={"parser": "text"},
        )

    def _parse_docx(self, path: Path, filename: str) -> ParsedDocument:
        if DocxDocument is None:
            raise ImportError("Install python-docx to parse DOCX files.")

        document = DocxDocument(path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return ParsedDocument(
            filename=filename,
            mime_type=(
                mimetypes.guess_type(filename)[0]
                or "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            pages=[ParsedPage(page_number=1, text=text.strip())],
            metadata={"parser": "python-docx"},
        )

    def _needs_ocr(self, text: str) -> bool:
        normalized = text.strip()
        if not normalized:
            return True
        return (len(normalized) / max(len(text), 1)) < self.ocr_min_text_threshold

    def _try_ocr_page(self, page: Any) -> str:
        try:
            ocr = self._get_ocr_engine()
            pixmap = page.get_pixmap()
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as image_file:
                image_path = image_file.name
            try:
                pixmap.save(image_path)
                result = ocr.ocr(image_path, cls=True)
            finally:
                if os.path.exists(image_path):
                    os.remove(image_path)
        except Exception:
            return ""

        lines = []
        for block in result or []:
            for line in block or []:
                if len(line) >= 2 and line[1]:
                    lines.append(str(line[1][0]))
        return "\n".join(lines)

    def _get_ocr_engine(self) -> Any:
        if self._ocr_engine is None:
            from paddleocr import PaddleOCR

            self._ocr_engine = PaddleOCR(lang=settings.OCR_LANGUAGE, show_log=False)
        return self._ocr_engine